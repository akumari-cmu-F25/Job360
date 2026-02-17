"""Resume parser for PDF and DOCX files."""

import logging
from typing import Optional, Dict, Any
from pathlib import Path
import docx
import PyPDF2
import pdfplumber

logger = logging.getLogger(__name__)


class ResumeParser:
    """Parses resume files (PDF/DOCX) and extracts text."""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.doc']
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        """
        Parse resume file and extract text.
        
        Args:
            file_path: Path to resume file
        
        Returns:
            Dict with 'text', 'metadata', 'format'
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"Resume file not found: {file_path}")
        
        file_ext = path.suffix.lower()
        
        if file_ext not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_ext}. Supported: {self.supported_formats}")
        
        logger.info(f"Parsing resume: {file_path} (format: {file_ext})")
        
        if file_ext == '.pdf':
            return self._parse_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            return self._parse_docx(file_path)
        else:
            raise ValueError(f"Unsupported format: {file_ext}")
    
    def _parse_pdf(self, file_path: str) -> Dict[str, Any]:
        """Parse PDF file."""
        text_parts = []
        metadata = {}
        
        try:
            # Try pdfplumber first (better for text extraction)
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                
                # Extract metadata
                if pdf.metadata:
                    metadata = {
                        'title': pdf.metadata.get('Title'),
                        'author': pdf.metadata.get('Author'),
                        'subject': pdf.metadata.get('Subject'),
                        'creator': pdf.metadata.get('Creator'),
                    }
        except Exception as e:
            logger.warning(f"pdfplumber failed, trying PyPDF2: {e}")
            # Fallback to PyPDF2
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                
                if pdf_reader.metadata:
                    metadata = {
                        'title': pdf_reader.metadata.get('/Title'),
                        'author': pdf_reader.metadata.get('/Author'),
                        'subject': pdf_reader.metadata.get('/Subject'),
                    }
        
        full_text = '\n\n'.join(text_parts)
        
        return {
            'text': full_text,
            'metadata': metadata,
            'format': 'pdf',
            'page_count': len(text_parts)
        }
    
    def _parse_docx(self, file_path: str) -> Dict[str, Any]:
        """Parse DOCX file."""
        try:
            doc = docx.Document(file_path)
            
            # Extract text from paragraphs
            text_parts = [para.text for para in doc.paragraphs if para.text.strip()]
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        text_parts.append(row_text)
            
            full_text = '\n\n'.join(text_parts)
            
            # Extract metadata
            core_props = doc.core_properties
            metadata = {
                'title': core_props.title,
                'author': core_props.author,
                'subject': core_props.subject,
                'created': str(core_props.created) if core_props.created else None,
            }
            
            return {
                'text': full_text,
                'metadata': metadata,
                'format': 'docx',
                'paragraph_count': len(text_parts)
            }
            
        except Exception as e:
            logger.error(f"Failed to parse DOCX: {e}")
            raise
