"""Resume exporter - Converts Profile to DOCX/PDF."""

import logging
from pathlib import Path
from typing import Optional
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.agents.profile_parser.profile_models import Profile

logger = logging.getLogger(__name__)


class ResumeExporter:
    """Exports Profile to DOCX format."""
    
    def __init__(self):
        pass
    
    def export_to_docx(self, profile: Profile, output_path: str) -> str:
        """
        Export profile to DOCX file.
        
        Args:
            profile: Profile to export
            output_path: Path to save DOCX file
        
        Returns:
            Path to saved file
        """
        doc = docx.Document()
        
        # Set up styles
        self._setup_styles(doc)
        
        # Header
        if profile.name:
            heading = doc.add_heading(profile.name, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contact information
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_para.style = 'Normal'
        
        contact_parts = []
        if profile.email:
            contact_parts.append(profile.email)
        if profile.phone:
            contact_parts.append(profile.phone)
        if profile.location:
            contact_parts.append(profile.location)
        if contact_parts:
            contact_para.add_run(" | ".join(contact_parts))
        
        if profile.linkedin or profile.github:
            links_para = doc.add_paragraph()
            links_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            links = []
            if profile.linkedin:
                links.append(f"LinkedIn: {profile.linkedin}")
            if profile.github:
                links.append(f"GitHub: {profile.github}")
            if profile.website:
                links.append(f"Website: {profile.website}")
            if links:
                links_para.add_run(" | ".join(links))
        
        doc.add_paragraph()  # Spacing
        
        # Professional Summary
        if profile.summary:
            doc.add_heading("Professional Summary", level=2)
            summary_para = doc.add_paragraph(profile.summary)
            summary_para.style = 'Normal'
            doc.add_paragraph()  # Spacing
        
        # Experience
        if profile.experiences:
            doc.add_heading("Experience", level=2)
            for exp in profile.experiences:
                # Title and company
                exp_para = doc.add_paragraph()
                exp_para.add_run(exp.title).bold = True
                exp_para.add_run(f" | {exp.company}")
                if exp.location:
                    exp_para.add_run(f", {exp.location}")
                
                # Dates
                if exp.start_date or exp.end_date:
                    date_para = doc.add_paragraph()
                    start = exp.start_date or ""
                    end = exp.end_date or "Present"
                    date_para.add_run(f"{start} - {end}")
                    date_para.style = 'Normal'
                
                # Bullets
                if exp.bullets:
                    for bullet in exp.bullets:
                        bullet_para = doc.add_paragraph(bullet, style='List Bullet')
                
                # Technologies
                if exp.technologies:
                    tech_para = doc.add_paragraph()
                    tech_para.add_run(f"Technologies: {', '.join(exp.technologies)}").italic = True
                
                doc.add_paragraph()  # Spacing between experiences
        
        # Education
        if profile.education:
            doc.add_heading("Education", level=2)
            for edu in profile.education:
                edu_para = doc.add_paragraph()
                edu_para.add_run(edu.degree).bold = True
                if edu.field_of_study:
                    edu_para.add_run(f" in {edu.field_of_study}")
                edu_para.add_run(f" | {edu.institution}")
                if edu.location:
                    edu_para.add_run(f", {edu.location}")
                
                if edu.graduation_date:
                    grad_para = doc.add_paragraph()
                    grad_para.add_run(f"Graduated: {edu.graduation_date}")
                    if edu.gpa:
                        grad_para.add_run(f" | GPA: {edu.gpa}")
                
                doc.add_paragraph()  # Spacing
        
        # Skills
        if profile.skills:
            doc.add_heading("Skills", level=2)
            skill_names = [s.name for s in profile.skills]
            skills_para = doc.add_paragraph(", ".join(skill_names))
            doc.add_paragraph()  # Spacing
        
        # Projects
        if profile.projects:
            doc.add_heading("Projects", level=2)
            for proj in profile.projects:
                proj_para = doc.add_paragraph()
                proj_para.add_run(proj.name).bold = True
                
                if proj.description:
                    desc_para = doc.add_paragraph(proj.description)
                
                if proj.technologies:
                    tech_para = doc.add_paragraph()
                    tech_para.add_run(f"Technologies: {', '.join(proj.technologies)}").italic = True
                
                doc.add_paragraph()  # Spacing
        
        # Save document
        doc.save(output_path)
        logger.info(f"Resume exported to {output_path}")
        
        return output_path
    
    def _setup_styles(self, doc):
        """Set up document styles."""
        styles = doc.styles
        
        # Normal style
        normal_style = styles['Normal']
        normal_font = normal_style.font
        normal_font.name = 'Calibri'
        normal_font.size = Pt(11)
        
        # Heading styles
        for i in range(1, 3):
            heading_style = styles[f'Heading {i}']
            heading_font = heading_style.font
            heading_font.name = 'Calibri'
            if i == 1:
                heading_font.size = Pt(16)
                heading_font.bold = True
            elif i == 2:
                heading_font.size = Pt(12)
                heading_font.bold = True
