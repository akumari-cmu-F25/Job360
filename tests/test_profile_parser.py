"""Tests for Profile Parser Agent."""

import pytest
import json
from unittest.mock import Mock, AsyncMock, patch
from pathlib import Path
import tempfile
import os

from src.agents.profile_parser import ProfileParserAgent, Profile
from src.orchestrator.central_orchestrator import CentralOrchestrator


@pytest.fixture
def orchestrator():
    """Create a test orchestrator."""
    return CentralOrchestrator()


@pytest.fixture
def profile_agent(orchestrator):
    """Create a test profile parser agent."""
    return ProfileParserAgent(orchestrator)


@pytest.fixture
def sample_resume_text():
    """Sample resume text for testing."""
    return """
John Doe
Email: john.doe@email.com | Phone: (555) 123-4567
Location: San Francisco, CA
LinkedIn: linkedin.com/in/johndoe | GitHub: github.com/johndoe

PROFESSIONAL SUMMARY
Experienced software engineer with expertise in distributed systems and machine learning.

EXPERIENCE

Software Engineer | Tech Corp | San Francisco, CA
Jan 2020 - Present
• Designed and implemented distributed systems using Python and PyTorch
• Led team of 5 engineers in developing ML models
• Technologies: Python, PyTorch, LLMs, LangChain, AWS

Intern | Startup Inc | San Francisco, CA
Jun 2019 - Dec 2019
• Developed web applications using React and Node.js
• Technologies: JavaScript, React, Node.js

EDUCATION

Bachelor of Science in Computer Science
University of California, Berkeley | Berkeley, CA
Graduated: May 2019
GPA: 3.8/4.0

SKILLS
Programming Languages: Python, JavaScript, TypeScript
Frameworks: React, PyTorch, LangChain
Cloud: AWS, GCP
"""


def test_profile_agent_initialization(profile_agent):
    """Test that profile agent initializes correctly."""
    assert profile_agent.name == "ProfileParserAgent"
    assert profile_agent.parser is not None
    assert profile_agent.normalizer is not None


def test_tech_normalizer():
    """Test technology normalization."""
    from src.agents.profile_parser.tech_normalizer import TechNormalizer
    
    normalizer = TechNormalizer()
    
    # Test direct mapping
    normalized, category = normalizer.normalize("python")
    assert normalized == "Python"
    
    normalized, category = normalizer.normalize("pytorch")
    assert normalized == "PyTorch"
    
    normalized, category = normalizer.normalize("llm")
    assert normalized == "LLMs"
    
    # Test case variations
    normalized, category = normalizer.normalize("PYTHON")
    assert normalized == "Python"
    
    normalized, category = normalizer.normalize("LangChain")
    assert normalized == "LangChain"


def test_resume_parser_docx():
    """Test DOCX parsing."""
    from src.agents.profile_parser.resume_parser import ResumeParser
    
    parser = ResumeParser()
    
    # Create a temporary DOCX file
    import docx
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    temp_path = temp_file.name
    temp_file.close()
    
    try:
        # Create a simple docx
        doc = docx.Document()
        doc.add_paragraph("John Doe")
        doc.add_paragraph("Software Engineer")
        doc.save(temp_path)
        
        # Parse it
        result = parser.parse(temp_path)
        
        assert "text" in result
        assert "format" in result
        assert result["format"] == "docx"
        assert "John Doe" in result["text"]
        
    finally:
        if os.path.exists(temp_path):
            os.unlink(temp_path)


@pytest.mark.asyncio
async def test_profile_parser_extract_structured_data(profile_agent, sample_resume_text):
    """Test structured data extraction."""
    with patch.object(profile_agent.client.chat.completions, 'create') as mock_create:
        mock_response = Mock()
        mock_response.choices = [Mock()]
        mock_response.choices[0].message.content = json.dumps({
            "personal_info": {
                "name": "John Doe",
                "email": "john.doe@email.com"
            },
            "summary": "Experienced software engineer",
            "experiences": [
                {
                    "title": "Software Engineer",
                    "company": "Tech Corp",
                    "start_date": "2020-01",
                    "end_date": "Present",
                    "is_current": True,
                    "bullets": ["Designed systems"],
                    "technologies": ["Python", "PyTorch"]
                }
            ],
            "education": [],
            "skills": [],
            "projects": [],
            "certifications": [],
            "awards": [],
            "languages": []
        })
        mock_create.return_value = mock_response
        
        result = await profile_agent._extract_structured_data(sample_resume_text)
        
        assert "personal_info" in result
        assert "experiences" in result
        assert len(result["experiences"]) > 0


@pytest.mark.asyncio
async def test_profile_parser_execute_with_invalid_file(profile_agent):
    """Test profile parser with invalid file path."""
    result = await profile_agent.execute(
        "nonexistent_file.pdf",
        input_type="resume_file"
    )
    
    assert result["success"] is False
    assert "violations" in result
